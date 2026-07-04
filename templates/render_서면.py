#!/usr/bin/env python3
"""docxtpl 템플릿 → 서면 docx 렌더링.

사용법:
    pip install docxtpl
    python render_서면.py <템플릿.docx> <context.json> <출력.docx>

예시:
    python render_서면.py 소장_템플릿_docxtpl.docx context_소장_예시.json 소장_샘플.docx
    python render_서면.py 준비서면_템플릿_docxtpl.docx context_준비서면_예시.json 준비서면_샘플.docx

연계:
    작성 스킬(draft-complaint/draft-brief)이 리서치/쟁점별_법리.md의 검증된 법리를
    context JSON의 청구원인(또는 본문) 배열에 채워 넣은 뒤 이 스크립트를 실행하면
    법원 양식 그대로의 워드 초안이 나온다. 실행 전 check_projection.py로 필수 키를 검사할 것.

검토 표시 후처리:
    렌더 직후 본문을 훑어 변호사 검토 지점을 색으로 표시한다.
    - [변호사 확인 필요…]  → 파란색 (수정·확정해야 할 값)
    - [출처: …], [행위시 확인: …] → 빨간색 (핀 블록 — 최종 제출본에서 변호사가 제거)
"""
import copy
import json
import re
import sys

from docx.shared import RGBColor
from docx.text.run import Run
from docxtpl import DocxTemplate

BLUE = RGBColor(0x00, 0x00, 0xFF)   # [변호사 확인 필요…]
RED = RGBColor(0xFF, 0x00, 0x00)    # [출처: …]

MARK_RE = re.compile(r"(\[변호사 확인 필요[^\]]*\]|\[출처:[^\]]*\]|\[행위시 확인[^\]]*\])")


def _set_run_text(r_element, text: str) -> None:
    """런 XML 요소의 텍스트를 통째로 교체한다(w:t 단일화, 공백 보존)."""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for t in r_element.findall(f"{ns}t"):
        r_element.remove(t)
    t = r_element.makeelement(f"{ns}t", {})
    t.text = text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    r_element.append(t)


def _colorize_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        parts = MARK_RE.split(run.text)
        if len(parts) <= 1:
            continue
        anchor = run._r
        for part in parts:
            if not part:
                continue
            new_r = copy.deepcopy(run._r)
            _set_run_text(new_r, part)
            new_run = Run(new_r, run._parent)
            if part.startswith("[변호사 확인 필요"):
                new_run.font.color.rgb = BLUE
            elif part.startswith("[출처:"):
                new_run.font.color.rgb = RED
            elif part.startswith("[행위시 확인"):
                new_run.font.color.rgb = RED
            anchor.addnext(new_r)
            anchor = new_r
        run._r.getparent().remove(run._r)


def colorize_review_marks(document) -> int:
    """본문·표 안의 검토 표시를 착색. 처리한 단락 수를 반환."""
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    n = 0
    for p in paragraphs:
        if MARK_RE.search(p.text):
            _colorize_paragraph(p)
            n += 1
    return n


def main() -> None:
    if len(sys.argv) != 4:
        sys.exit(__doc__)
    template_path, context_path, output_path = sys.argv[1:4]

    with open(context_path, encoding="utf-8") as f:
        context = json.load(f)

    doc = DocxTemplate(template_path)
    # autoescape 필수: 값에 '<' '>' '&'가 있으면(예: 조문의 "<신설 2014.3.24>",
    # 연혁 핀의 "부칙 <제20520호,...>") XML이 깨져 Word가 파일을 열지 못한다.
    doc.render(context, autoescape=True)
    marked = colorize_review_marks(doc.docx)
    doc.save(output_path)
    print(f"렌더링 완료: {output_path} (검토 표시 착색 단락 {marked}개)")


if __name__ == "__main__":
    main()
